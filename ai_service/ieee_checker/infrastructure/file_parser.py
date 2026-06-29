from __future__ import annotations
import logging
import os
from typing import Tuple

import pdfplumber
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

logger = logging.getLogger(__name__)


def extract_text_from_pdf(pdf_path: str) -> Tuple[str, int]:
    pages_text = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
        return "\n".join(pages_text), len(pages_text)
    except Exception as e:
        logger.error("PDF extraction failed: %s", e)
        return "", 0


def extract_text_from_docx(docx_path: str) -> Tuple[str, int]:
    if not DOCX_AVAILABLE:
        logger.error("python-docx not installed. Cannot process DOCX files.")
        return "", 0
    try:
        doc = Document(docx_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        return "\n".join(full_text), 1
    except Exception as e:
        logger.error("DOCX extraction failed: %s", e)
        return "", 0


def extract_text_from_file(file_path: str) -> Tuple[str, int, str]:
    file_ext = os.path.splitext(file_path)[1].lower()
    if file_ext == '.pdf':
        text, pages = extract_text_from_pdf(file_path)
        return text, pages, 'pdf'
    elif file_ext == '.docx':
        text, pages = extract_text_from_docx(file_path)
        return text, pages, 'docx'
    logger.error("Unsupported file type: %s", file_ext)
    return "", 0, 'unknown'
