from __future__ import annotations
import re
import time
import logging
import os
from dataclasses import dataclass, field, asdict
from typing import Any, Dict, List, Optional, Tuple

import pdfplumber
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None  

logger = logging.getLogger(__name__)


@dataclass
class ReferenceEntry:
    index: int
    raw_text: str
    authors: str = ""
    title: str = ""
    source: str = ""
    year: str = ""
    doi: str = ""
    url: str = ""
    volume: str = ""
    issue: str = ""
    pages: str = ""
    format_errors: List[str] = field(default_factory=list)
    crossref_verified: Optional[bool] = None
    crossref_message: str = ""
    ieee_score: float = 0.0


@dataclass
class IEEECheckResult:
    paper_title: str = ""
    detected_language: str = "unknown"
    total_pages: int = 0
    citations_in_text: List[int] = field(default_factory=list)
    citations_missing_from_references: List[int] = field(default_factory=list)
    unused_references: List[int] = field(default_factory=list)
    references: List[ReferenceEntry] = field(default_factory=list)
    total_references: int = 0
    format_issues_summary: List[str] = field(default_factory=list)
    references_without_year: List[int] = field(default_factory=list)
    references_without_authors: List[int] = field(default_factory=list)
    references_without_title: List[int] = field(default_factory=list)
    crossref_checked: int = 0
    crossref_verified_count: int = 0
    crossref_failed: List[int] = field(default_factory=list)
    citation_matching_score: float = 0.0
    format_score: float = 0.0
    crossref_score: float = 0.0
    overall_score: float = 0.0
    summary: str = ""
    recommendations: List[str] = field(default_factory=list)
    status: str = "pending"

    def to_dict(self) -> Dict[str, Any]:
        d = asdict(self)
        return d

def extract_text_from_pdf(pdf_path: str) -> Tuple[str, int]:
    pages_text: List[str] = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
        return "\n".join(pages_text), len(pages_text)
    except Exception as e:
        logger.error("PDF extraction failed: %s", e)
        return "", 0


def extract_text_from_docx(docx_path: str) -> Tuple[str, int]:
    if not DOCX_AVAILABLE:
        logger.error("python-docx not installed. Cannot process DOCX files.")
        return "", 0
    
    try:
        doc = Document(docx_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return "\n".join(full_text), 1 
    except Exception as e:
        logger.error("DOCX extraction failed: %s", e)
        return "", 0


def extract_text_from_file(file_path: str) -> Tuple[str, int, str]:
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        text, page_count = extract_text_from_pdf(file_path)
        return text, page_count, 'pdf'
    elif file_ext == '.docx':
        text, page_count = extract_text_from_docx(file_path)
        return text, page_count, 'docx'
    else:
        logger.error("Unsupported file type: %s", file_ext)
        return "", 0, 'unknown'


def detect_language(text: str) -> str:
    arabic_chars = len(re.findall(r'[\u0600-\u06FF]', text))
    latin_chars = len(re.findall(r'[a-zA-Z]', text))
    total = arabic_chars + latin_chars
    if total == 0:
        return "unknown"
    ar_ratio = arabic_chars / total
    if ar_ratio > 0.6:
        return "ar"
    if ar_ratio < 0.2:
        return "en"
    return "mixed"


def extract_paper_title(full_text: str) -> str:
    lines = [l.strip() for l in full_text.split('\n') if l.strip()]
    skip_kw = {
        'abstract', 'introduction', 'keywords', 'references',
        'ملخص', 'مقدمة', 'كلمات مفتاحية', 'المراجع',
    }
    for line in lines[:20]:
        if len(line) < 10:
            continue
        if any(kw in line.lower() for kw in skip_kw):
            continue
        if re.match(r'^\d+$', line):
            continue
        if len(line) > 15:
            return line
    return lines[0] if lines else ""

def extract_in_text_citations(full_text: str) -> Dict[int, int]:
 
    citation_counts: Dict[int, int] = {}

    for match in re.findall(r'\[(\d+(?:\s*,\s*\d+)*)\]', full_text):
        for num_str in re.findall(r'\d+', match):
            num = int(num_str)
            citation_counts[num] = citation_counts.get(num, 0) + 1

    for start_s, end_s in re.findall(r'\[(\d+)\]\s*[-\u2013]\s*\[(\d+)\]', full_text):
        for num in range(int(start_s), int(end_s) + 1):
            citation_counts[num] = citation_counts.get(num, 0) + 1

    return citation_counts



_REF_SECTION_PATTERNS = [
    r'\bReferences\b',
    r'\bBibliography\b',
    r'\bWorks Cited\b',
    r'\u0627\u0644\u0645\u0631\u0627\u062c\u0639',       
    r'\u0642\u0627\u0626\u0645\u0629 \u0627\u0644\u0645\u0635\u0627\u062f\u0631',  
]


def find_references_section(full_text: str) -> str:
    pattern = '|'.join(_REF_SECTION_PATTERNS)
    match = re.search(pattern, full_text, re.IGNORECASE)
    if not match:
        return ""
    return full_text[match.start():]


def split_references_section(ref_section: str) -> List[Tuple[int, str]]:
    results: List[Tuple[int, str]] = []
    ref_pattern = re.compile(
        r'(?:\[(\d+)\]\.?\s*|(?<!\d)(\d+)\.\s+)(.*?)(?=\n\s*(?:\[\d+\]\.?\s*|\d+\.\s+)|\Z)',
        re.DOTALL,
    )
    for m in ref_pattern.finditer(ref_section):
        num_str = m.group(1) or m.group(2)
        text = m.group(3).strip()
        if num_str and text:
            clean_text = re.sub(r'\s+', ' ', text).strip()
            if len(clean_text) > 5:
                results.append((int(num_str), clean_text))
    return results


_YEAR_RE = re.compile(r'\b(19|20)\d{2}\b')
_DOI_RE = re.compile(r'(?:doi\.org/|DOI:\s*|doi:\s*)(10\.\d{4,}/[^\s,\]]+)', re.IGNORECASE)
_URL_RE = re.compile(r'https?://[^\s\]]+', re.IGNORECASE)
_PAGES_RE = re.compile(r'pp?\.\s*(\d+[\s\-\u2013]+\d+)', re.IGNORECASE)
_VOL_RE = re.compile(r'vol\.?\s*(\d+)', re.IGNORECASE)
_NO_RE = re.compile(r'no\.?\s*(\d+)', re.IGNORECASE)


def parse_ieee_reference(ref_text: str) -> Dict[str, str]:
    result = {
        'authors': '', 'title': '', 'source': '',
        'year': '', 'doi': '', 'url': '',
        'volume': '', 'issue': '', 'pages': '',
    }
    text = ref_text.strip()

    doi_m = _DOI_RE.search(text)
    if doi_m:
        result['doi'] = doi_m.group(1)

    url_m = _URL_RE.search(text)
    if url_m:
        result['url'] = url_m.group(0)

    years = _YEAR_RE.findall(text)
    if years:
        result['year'] = years[-1]

    pages_m = _PAGES_RE.search(text)
    if pages_m:
        result['pages'] = pages_m.group(1)

    vol_m = _VOL_RE.search(text)
    if vol_m:
        result['volume'] = vol_m.group(1)

    no_m = _NO_RE.search(text)
    if no_m:
        result['issue'] = no_m.group(1)

    title_m = re.search(r'["\u201C\u201D\u0022]([^"\u201C\u201D\u0022]+)["\u201C\u201D\u0022]', text)
    if title_m:
        result['title'] = title_m.group(1).strip()

    if result['title']:
        authors_part = text[:text.find(result['title'])].strip().rstrip(',"')
        result['authors'] = re.sub(r'\s+', ' ', authors_part).strip()
    else:
        auth_m = re.match(r'^([A-Z]\.\s+[A-Z][a-zA-Z\-]+(?:(?:,\s*|,?\s+and\s+)[A-Z]\.\s+[A-Z][a-zA-Z\-]+)*(?:,?\s*et al\.?)?)', text)
        if auth_m:
            result['authors'] = auth_m.group(1).strip()
        else:
            ar_m = re.match(r'^([\u0600-\u06FF\s\u060C،,]+(?:[\u060C،,]\s*))', text)
            if ar_m and len(ar_m.group(1).strip()) > 3:
                result['authors'] = ar_m.group(1).strip()

    if result['title'] and result['year']:
        ti = text.find(result['title'])
        yi = text.rfind(result['year'])
        if ti >= 0 and yi > ti:
            after = text[ti + len(result['title']):yi]
            after = re.sub(r',?\s*vol\.\s*\d+.*', '', after, flags=re.IGNORECASE)
            after = re.sub(r'[,."]+', '', after).strip()
            result['source'] = after[:100]

    return result


def validate_ieee_format(ref: ReferenceEntry) -> List[str]:
    errors: List[str] = []

    if not ref.authors:
        errors.append("مفقود: اسم المؤلف/المؤلفين")
    else:
        arabic_count = len(re.findall(r'[\u0600-\u06FF]', ref.authors))
        if arabic_count == 0 and not re.search(r'[A-Z]\.\s*[A-Z][a-z]', ref.authors):
            errors.append("تنسيق المؤلفين: يُفضَّل النمط IEEE (الحرف.الأول الاسمُ الأخير)")

    if not ref.title:
        errors.append('مفقود: عنوان البحث (يجب بين علامتي اقتباس "...")')

    if not ref.year:
        errors.append("مفقودة: سنة النشر")
    elif not re.match(r'^(19|20)\d{2}$', ref.year):
        errors.append(f"سنة النشر غير صحيحة: {ref.year}")

    if not ref.source and not ref.url:
        errors.append("مفقود: اسم المجلة أو المؤتمر أو الناشر")

    if ref.doi and not re.match(r'^10\.\d{4,}/', ref.doi):
        errors.append(f"تنسيق DOI غير صحيح: {ref.doi}")

    if not ref.pages and not ref.url and ref.source:
        errors.append("يُنصح: إضافة أرقام الصفحات (pp. X-Y)")

    return errors

def verify_doi_crossref(doi: str) -> Tuple[bool, str]:
    try:
        import urllib.request
        import json

        url = f"https://api.crossref.org/works/{doi}"
        req = urllib.request.Request(
            url,
            headers={"User-Agent": "ASPU-Insight/1.0 (mailto:aspu@aspu.edu.sy)"}
        )
        with urllib.request.urlopen(req, timeout=5) as resp:
            data = json.loads(resp.read().decode())
            if data.get('status') == 'ok':
                titles = data.get('message', {}).get('title', [])
                t = titles[0][:80] if titles else "N/A"
                return True, f"تم التحقق ✓ العنوان في Crossref: {t}"
    except Exception as e:
        return False, f"تعذّر التحقق عبر Crossref: {str(e)[:60]}"
    return False, "لم يُوجد DOI في Crossref"


def calculate_scores(result: IEEECheckResult) -> IEEECheckResult:
    total_cited = len(result.citations_in_text)
    missing = len(result.citations_missing_from_references)
    unused = len(result.unused_references)

    if total_cited == 0 and result.total_references == 0:
        result.citation_matching_score = 100.0
    elif total_cited == 0:
        result.citation_matching_score = 50.0
    else:
        penalty = (missing * 10) + (unused * 5)
        result.citation_matching_score = max(0.0, 100.0 - penalty)

    if result.total_references == 0:
        result.format_score = 0.0
    else:
        total_errors = sum(len(ref.format_errors) for ref in result.references)
        avg_errors = total_errors / result.total_references
        result.format_score = max(0.0, 100.0 - (avg_errors * 15))

    if result.crossref_checked == 0:
        result.crossref_score = 100.0
    else:
        result.crossref_score = (result.crossref_verified_count / result.crossref_checked) * 100.0

    result.overall_score = round(
        (result.citation_matching_score * 0.40)
        + (result.format_score * 0.40)
        + (result.crossref_score * 0.20),
        1,
    )

    if result.overall_score >= 75:
        result.status = "pass"
    elif result.overall_score >= 50:
        result.status = "warning"
    else:
        result.status = "fail"

    return result


def generate_recommendations(result: IEEECheckResult) -> List[str]:
    recs: List[str] = []

    if result.citations_missing_from_references:
        nums = ", ".join(f"[{n}]" for n in sorted(result.citations_missing_from_references))
        recs.append(f"أضف المراجع المفقودة لهذه الاستشهادات: {nums}")

    if result.unused_references:
        nums = ", ".join(f"[{n}]" for n in sorted(result.unused_references))
        recs.append(f"المراجع غير المُستشهَد بها: {nums} — احذفها أو أضف استشهادات لها في النص")

    if result.references_without_year:
        nums = ", ".join(f"[{n}]" for n in sorted(result.references_without_year)[:5])
        recs.append(f"أضف سنة النشر للمراجع: {nums}")

    if result.references_without_authors:
        nums = ", ".join(f"[{n}]" for n in sorted(result.references_without_authors)[:5])
        recs.append(f"أضف أسماء المؤلفين للمراجع: {nums}")

    if result.references_without_title:
        nums = ", ".join(f"[{n}]" for n in sorted(result.references_without_title)[:5])
        recs.append(f"أضف عنوان البحث بين علامتي اقتباس للمراجع: {nums}")

    if result.crossref_failed:
        nums = ", ".join(f"[{n}]" for n in sorted(result.crossref_failed)[:3])
        recs.append(f"تحقق يدوياً من صحة DOI في المراجع: {nums}")

    if result.format_score < 60:
        recs.append(
            'راجع تنسيق IEEE العام: المؤلف (الحرف.الأول الاسم)، '
            '"العنوان," المجلة، vol. X, no. Y, pp. Z-W, السنة.'
        )

    if not recs:
        recs.append("ممتاز! لا توجد مشكلات جوهرية في الاستشهادات والمراجع.")

    return recs


def generate_summary(result: IEEECheckResult) -> str:
    status_ar = {
        "pass": "مقبول",
        "warning": "يحتاج تحسين",
        "fail": "يحتاج مراجعة جوهرية",
    }.get(result.status, "غير محدد")

    return (
        f'الورقة: "{result.paper_title[:60]}" | '
        f"اللغة: {result.detected_language} | "
        f"الصفحات: {result.total_pages} | "
        f"الاستشهادات في النص: {len(result.citations_in_text)} | "
        f"إجمالي المراجع: {result.total_references} | "
        f"الدرجة: {result.overall_score}/100 | "
        f"الحالة: {status_ar}"
    )


def perform_ieee_analysis(
    file_path: str,
    verify_crossref: bool = True,
    max_crossref_calls: int = 5,
) -> Dict[str, Any]:

    result = IEEECheckResult()

    full_text, page_count, file_type = extract_text_from_file(file_path)
    result.total_pages = page_count

    if not full_text.strip():
        result.summary = "تعذّر استخراج النص من الملف. تأكد أن الـ PDF يحتوي على نص قابل للبحث."
        result.status = "fail"
        result.recommendations = ["تحويل الـ PDF إلى صيغة نصية أو استخدام OCR."]
        return result.to_dict()

    result.paper_title = extract_paper_title(full_text)
    result.detected_language = detect_language(full_text)

    citation_counts = extract_in_text_citations(full_text)
    result.citations_in_text = sorted(citation_counts.keys())

    ref_section = find_references_section(full_text)
    raw_refs = split_references_section(ref_section) if ref_section else []

    ref_map: Dict[int, ReferenceEntry] = {}
    crossref_calls = 0

    for ref_num, ref_text in raw_refs:
        parsed = parse_ieee_reference(ref_text)
        entry = ReferenceEntry(
            index=ref_num,
            raw_text=ref_text[:300],
            authors=parsed['authors'],
            title=parsed['title'],
            source=parsed['source'],
            year=parsed['year'],
            doi=parsed['doi'],
            url=parsed['url'],
            volume=parsed['volume'],
            issue=parsed['issue'],
            pages=parsed['pages'],
        )

        entry.format_errors = validate_ieee_format(entry)
        entry.ieee_score = max(0.0, 100.0 - len(entry.format_errors) * 20)

        if verify_crossref and entry.doi and crossref_calls < max_crossref_calls:
            is_valid, msg = verify_doi_crossref(entry.doi)
            entry.crossref_verified = is_valid
            entry.crossref_message = msg
            crossref_calls += 1
            result.crossref_checked += 1
            if is_valid:
                result.crossref_verified_count += 1
            else:
                result.crossref_failed.append(ref_num)
            time.sleep(0.3)

        ref_map[ref_num] = entry

    result.references = list(ref_map.values())
    result.total_references = len(ref_map)
    ref_ids = set(ref_map.keys())

    cited_set = set(result.citations_in_text)
    result.citations_missing_from_references = sorted(cited_set - ref_ids)
    result.unused_references = sorted(ref_ids - cited_set)

    for ref in result.references:
        joined_errors = " ".join(ref.format_errors)
        if "سنة" in joined_errors:
            result.references_without_year.append(ref.index)
        if "المؤلف" in joined_errors:
            result.references_without_authors.append(ref.index)
        if "عنوان" in joined_errors:
            result.references_without_title.append(ref.index)

    all_format_errors: List[str] = []
    for ref in result.references:
        for err in ref.format_errors:
            all_format_errors.append(f"[{ref.index}]: {err}")
    result.format_issues_summary = all_format_errors[:20]

    result = calculate_scores(result)

    result.recommendations = generate_recommendations(result)
    result.summary = generate_summary(result)

    return result.to_dict()
